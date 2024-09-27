import fitz  # PyMuPDF
import markdown
from docx import Document
import olefile
import pyhwp

def get_resume_samples(sample_file):
    """이력서 샘플 데이터를 불러오는 함수 (다양한 형식 지원)"""
    file_extension = sample_file.split('.')[-1].lower()

    try:
        if file_extension == "pdf":
            return read_pdf(sample_file)
        elif file_extension == "md":
            return read_markdown(sample_file)
        elif file_extension == "docx":
            return read_docx(sample_file)
        elif file_extension == "hwp":
            return read_hwp(sample_file)
        else:
            return "지원하지 않는 파일 형식입니다."
    except FileNotFoundError:
        return "Sample resume file not found. Please check the file path."
    except Exception as e:
        return f"Error while reading the file: {str(e)}"

def read_pdf(file_path):
    """PDF 파일 읽기"""
    resume_samples = ""
    with fitz.open(file_path) as pdf_file:
        for page_num in range(len(pdf_file)):
            page = pdf_file.load_page(page_num)
            resume_samples += page.get_text()
    return resume_samples

def read_markdown(file_path):
    """Markdown 파일 읽기"""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def read_docx(file_path):
    """DOCX 파일 읽기"""
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def read_hwp(file_path):
    """HWP 파일 읽기 (olefile 및 pyhwp로 시도)"""
    try:
        # 먼저 olefile로 시도
        hwp_file = olefile.OleFileIO(file_path)
        
        if hwp_file.exists('HwpSummaryInformation'):
            stream = hwp_file.openstream('HwpSummaryInformation')
            data = stream.read()
            
            # UTF-16으로 인코딩된 경우를 가정하여 디코딩
            text = data.decode('utf-16')
            return text

        return "HWP 파일에서 텍스트를 찾을 수 없습니다."

    except Exception as olefile_error:
        print(f"olefile을 사용한 HWP 파일 처리 실패: {olefile_error}")
        print("pyhwp로 재시도합니다...")

        # pyhwp로 재시도
        try:
            hwp_doc = pyhwp.HWPDocument(file_path)
            # pyhwp로 HWP 파일의 모든 텍스트 추출
            text = ""
            for section in hwp_doc.sections:
                for paragraph in section.paragraphs:
                    text += paragraph.text + "\n"
            return text
        except Exception as pyhwp_error:
            return f"olefile과 pyhwp 모두 실패했습니다. 다른 파일 형식을 사용해 주세요: {pyhwp_error}"